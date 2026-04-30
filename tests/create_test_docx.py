#!/usr/bin/env python3
"""Create test exam DOCX file for Study Mode validation"""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document

def create_exam_docx():
    doc = Document()
    doc.add_heading('DE THI THPT QUOC GIA 2025', 0)
    doc.add_paragraph('Mon: Toan')
    doc.add_paragraph('Thoi gian: 90 phut')
    doc.add_paragraph()

    # Câu 1
    doc.add_paragraph('Cau 1: Gia tri cua bieu thuc √(16) + |-3| la:')
    doc.add_paragraph('A. 1')
    doc.add_paragraph('B. 7')
    doc.add_paragraph('C. -1')
    doc.add_paragraph('D. -7')
    doc.add_paragraph()

    # Câu 2
    doc.add_paragraph('Cau 2: Phuong trinh x² - 5x + 6 = 0 co nghiem la:')
    doc.add_paragraph('A. x₁=1, x₂=6')
    doc.add_paragraph('B. x₁=2, x₂=3')
    doc.add_paragraph('C. x₁=-2, x₂=-3')
    doc.add_paragraph('D. x₁=3, x₂=2')
    doc.add_paragraph()

    # Câu 3
    doc.add_paragraph('Cau 3: Tinh lim(x→2) (x²-4)/(x-2):')
    doc.add_paragraph('A. 0')
    doc.add_paragraph('B. 2')
    doc.add_paragraph('C. 4')
    doc.add_paragraph('D. vo cung')
    doc.add_paragraph()

    # Câu 4
    doc.add_paragraph('Cau 4: Ham so y = sin(x) co chu ky:')
    doc.add_paragraph('A. π')
    doc.add_paragraph('B. 2π')
    doc.add_paragraph('C. 90°')
    doc.add_paragraph('D. 180°')
    doc.add_paragraph()

    # Câu 5
    doc.add_paragraph('Cau 5: Tich phan ∫₀¹ x dx bang:')
    doc.add_paragraph('A. 0')
    doc.add_paragraph('B. 0.5')
    doc.add_paragraph('C. 1')
    doc.add_paragraph('D. 2')

    filename = "test_exam.docx"
    doc.save(filename)
    print(f"Created {filename} ({os.path.getsize(filename)} bytes)")
    return filename

if __name__ == "__main__":
    fname = create_exam_docx()
    print(f"File location: {os.path.abspath(fname)}")
