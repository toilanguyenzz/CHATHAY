"""Document parser: extracts text from PDF, Word, and images."""

import os
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


async def parse_pdf(file_path: str, max_pages: int = 100) -> str:
    """Extract text from PDF using pdfplumber (fallback to PyMuPDF)."""
    text_parts = []

    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            pages_to_read = min(total_pages, max_pages)

            for i in range(pages_to_read):
                page = pdf.pages[i]
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            if text_parts:
                result = "\n\n".join(text_parts)
                logger.info(f"PDF parsed with pdfplumber: {pages_to_read}/{total_pages} pages, {len(result)} chars")
                return result

    except Exception as e:
        logger.warning(f"pdfplumber failed, trying PyMuPDF: {e}")

    # Fallback to PyMuPDF
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        total_pages = len(doc)
        pages_to_read = min(total_pages, max_pages)

        for i in range(pages_to_read):
            page = doc[i]
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)

        doc.close()

        if text_parts:
            result = "\n\n".join(text_parts)
            logger.info(f"PDF parsed with PyMuPDF: {pages_to_read}/{total_pages} pages, {len(result)} chars")
            return result

    except Exception as e:
        logger.error(f"PyMuPDF also failed: {e}")

    return ""


async def parse_docx(file_path: str) -> str:
    """Extract text from Word .docx file."""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        result = "\n".join(text_parts)
        logger.info(f"DOCX parsed: {len(text_parts)} paragraphs, {len(result)} chars")
        return result

    except Exception as e:
        logger.warning(f"python-docx failed: {e} — trying raw XML fallback...")

    # ── FALLBACK: Read DOCX as ZIP → extract raw XML text ──
    # Handles non-standard DOCX from Google Docs, WPS Office, etc.
    try:
        import zipfile
        import re

        with zipfile.ZipFile(file_path, 'r') as z:
            # Try standard path first, then scan for document XML
            xml_content = None
            doc_paths = ['word/document.xml', 'word/document2.xml']

            for dp in doc_paths:
                if dp in z.namelist():
                    xml_content = z.read(dp).decode('utf-8', errors='ignore')
                    break

            # If standard paths don't work, find any XML with body content
            if not xml_content:
                for name in z.namelist():
                    if name.endswith('.xml') and 'document' in name.lower():
                        xml_content = z.read(name).decode('utf-8', errors='ignore')
                        break

            if not xml_content:
                logger.error("DOCX fallback: no document XML found in ZIP")
                return ""

            # Strip XML tags, keep text content
            # Match <w:t ...>text</w:t> tags (Word text runs)
            text_runs = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
            if not text_runs:
                # Broader fallback: strip all XML tags
                raw_text = re.sub(r'<[^>]+>', ' ', xml_content)
                raw_text = re.sub(r'\s+', ' ', raw_text).strip()
                if len(raw_text) > 50:
                    logger.info(f"DOCX fallback (broad strip): {len(raw_text)} chars")
                    return raw_text
                return ""

            # Join text runs, use paragraph breaks where appropriate
            # <w:p> tags indicate paragraph boundaries
            paragraphs = re.split(r'<w:p[ >]', xml_content)
            text_parts = []
            for para in paragraphs:
                runs = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', para)
                if runs:
                    para_text = ''.join(runs).strip()
                    if para_text:
                        text_parts.append(para_text)

            result = "\n".join(text_parts)
            logger.info(f"DOCX fallback (XML): {len(text_parts)} paragraphs, {len(result)} chars")
            return result

    except Exception as fallback_err:
        logger.error(f"DOCX fallback also failed: {fallback_err}")
        return ""


async def parse_xlsx(file_path: str, max_rows_per_sheet: int = 500) -> str:
    """Extract text from Excel .xlsx — tất cả sheets, format bảng."""
    try:
        from openpyxl import load_workbook

        # Load in read-only mode for performance
        wb = load_workbook(file_path, data_only=True, read_only=True)
        text_parts = []
        total_rows = 0

        for sheet in wb.worksheets:
            # Skip hidden sheets
            if sheet.sheet_state != 'visible':
                logger.info(f"Skipping hidden sheet: {sheet.title}")
                continue

            sheet_name = sheet.title
            max_row = sheet.max_row
            max_col = sheet.max_column

            if max_row == 0 or max_col == 0:
                continue  # Skip empty sheets

            # Header for this sheet
            text_parts.append(f"\n📋 Sheet: \"{sheet_name}\" ({min(max_row, max_rows_per_sheet)} hàng × {max_col} cột)")
            text_parts.append("─" * 60)

            # Read rows
            rows_read = 0
            for row in sheet.iter_rows(values_only=True, max_row=max_rows_per_sheet):
                rows_read += 1
                # Format row: join cells with " | "
                row_cells = []
                for cell in row:
                    if cell is None:
                        cell_text = ""
                    else:
                        cell_text = str(cell).strip()
                    row_cells.append(cell_text)

                row_line = " | ".join(row_cells)
                if row_line.strip():
                    text_parts.append(row_line)

            total_rows += rows_read
            text_parts.append("")  # blank line between sheets

        wb.close()

        result = "\n".join(text_parts)
        if result.strip():
            summary = f"[Excel: {len(wb.worksheets)} sheets, {total_rows} hàng tổng]"
            logger.info(f"Excel parsed: {len(wb.worksheets)} sheets, {total_rows} rows, {len(result)} chars")
            return result + "\n\n" + summary
        else:
            logger.warning("Excel file appears empty")
            return ""

    except Exception as e:
        logger.error(f"Excel parsing failed: {e}")
        return ""


async def convert_pdf_to_images(file_path: str, max_pages: int = 10, dpi: int = 200) -> list[str]:
    """Convert PDF pages to PNG images for OCR fallback.
    
    Used when pdfplumber/PyMuPDF can't extract text (scanned/handwritten PDFs).
    Returns list of image file paths. Caller is responsible for cleanup.
    """
    image_paths = []
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        total_pages = len(doc)
        pages_to_render = min(total_pages, max_pages)
        zoom = dpi / 72  # 72 is default PDF DPI
        matrix = fitz.Matrix(zoom, zoom)

        base_dir = os.path.dirname(file_path)
        base_name = os.path.splitext(os.path.basename(file_path))[0]

        for i in range(pages_to_render):
            page = doc[i]
            pix = page.get_pixmap(matrix=matrix)
            img_path = os.path.join(base_dir, f"{base_name}_page_{i+1}.png")
            pix.save(img_path)
            image_paths.append(img_path)
            logger.info(f"PDF page {i+1}/{pages_to_render} rendered to image: {pix.width}x{pix.height}")

        doc.close()
        logger.info(f"PDF converted to {len(image_paths)} images for OCR fallback")

    except Exception as e:
        logger.error(f"PDF to image conversion failed: {e}")
        # Cleanup any partially created files
        for path in image_paths:
            try:
                os.remove(path)
            except OSError:
                pass
        image_paths = []

    return image_paths


async def parse_image_ocr(file_path: str) -> Optional[str]:
    """
    Extract text from image using Gemini's vision capability.
    This avoids needing a separate Google Cloud Vision API key.
    Returns None if OCR should be handled by the AI summarizer directly.
    """
    # We return None here because Gemini can directly read images
    # The AI summarizer will handle image files directly via multimodal
    return None


def get_file_type(file_path: str) -> str:
    """Detect file type from extension."""
    ext = os.path.splitext(file_path)[1].lower()

    type_map = {
        ".pdf": "pdf",
        ".doc": "doc_legacy",
        ".docx": "docx",
        ".xlsx": "xlsx",
        ".xls": "xls_legacy",
        ".jpg": "image",
        ".jpeg": "image",
        ".png": "image",
        ".bmp": "image",
        ".webp": "image",
        ".tiff": "image",
        ".tif": "image",
    }

    return type_map.get(ext, "unknown")


async def parse_doc_legacy(file_path: str) -> str:
    """Extract text from legacy .doc (Word 97-2003) binary format."""

    # Method 1: Try olefile (OLE2 compound file reader)
    try:
        import olefile
        import re

        ole = olefile.OleFileIO(file_path)
        # The main text stream in .doc files
        if ole.exists('WordDocument'):
            # Try to read the text from the Word Document stream
            # .doc stores text in the 'WordDocument' stream in binary format
            # But readable text is often in '1Table' or '0Table' streams
            text_parts = []

            # Try reading all streams and extract readable text
            for stream_path in ole.listdir():
                stream_name = '/'.join(stream_path)
                try:
                    data = ole.openstream(stream_path).read()
                    # Try UTF-16LE decoding (Word .doc uses UTF-16LE for Unicode text)
                    try:
                        decoded = data.decode('utf-16-le', errors='ignore')
                        # Filter only printable text chunks
                        chunks = re.findall(r'[\w\s.,;:!?()\[\]{}\-–—"\'/+=%@#$&*~`^\\|<>àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ]{3,}', decoded)
                        for chunk in chunks:
                            clean = chunk.strip()
                            if len(clean) > 5 and not clean.startswith('\x00'):
                                text_parts.append(clean)
                    except Exception:
                        pass
                except Exception:
                    continue

            ole.close()

            if text_parts:
                # Deduplicate while preserving order
                seen = set()
                unique_parts = []
                for part in text_parts:
                    if part not in seen:
                        seen.add(part)
                        unique_parts.append(part)

                result = '\n'.join(unique_parts)
                if len(result) > 50:
                    logger.info(f"DOC legacy parsed (olefile): {len(unique_parts)} chunks, {len(result)} chars")
                    return result

        ole.close()
    except ImportError:
        logger.warning("olefile not installed, skipping OLE method")
    except Exception as e:
        logger.warning(f"olefile method failed: {e}")

    # Method 2: Raw binary text extraction (brute force but effective)
    try:
        import re

        with open(file_path, 'rb') as f:
            raw_data = f.read()

        text_parts = []

        # Try UTF-16LE (common in .doc)
        try:
            decoded_utf16 = raw_data.decode('utf-16-le', errors='ignore')
            # Extract meaningful text runs (at least 4 chars, Vietnamese-aware)
            runs = re.findall(
                r'[A-Za-z0-9àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ][A-Za-z0-9 .,;:!?()\-–\"\'/àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđÀÁẠẢÃÂẦẤẬẨẪĂẰẮẶẲẴÈÉẸẺẼÊỀẾỆỂỄÌÍỊỈĨÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠÙÚỤỦŨƯỪỨỰỬỮỲÝỴỶỸĐ]{4,}',
                decoded_utf16
            )
            text_parts.extend(runs)
        except Exception:
            pass

        # Try Latin-1 / CP1252 fallback (older .doc files)
        if not text_parts:
            try:
                decoded_latin = raw_data.decode('cp1252', errors='ignore')
                runs = re.findall(r'[\x20-\x7E\xC0-\xFF]{5,}', decoded_latin)
                text_parts.extend(runs)
            except Exception:
                pass

        if text_parts:
            # Filter noise: remove very short or binary-looking strings
            clean_parts = []
            for part in text_parts:
                stripped = part.strip()
                # Skip if mostly control chars or too short
                if len(stripped) < 5:
                    continue
                # Skip if more than 30% non-letter chars (likely binary noise)
                letter_count = sum(1 for c in stripped if c.isalpha() or c.isspace())
                if letter_count / len(stripped) < 0.5:
                    continue
                clean_parts.append(stripped)

            if clean_parts:
                result = '\n'.join(clean_parts)
                logger.info(f"DOC legacy parsed (binary): {len(clean_parts)} text runs, {len(result)} chars")
                return result

    except Exception as e:
        logger.error(f"DOC legacy binary extraction failed: {e}")

    logger.error(f"DOC legacy: all methods failed for {file_path}")
    return ""


async def extract_text(file_path: str, max_pages: int = 100) -> tuple[str, str]:
    """
    Extract text from a document file.

    Returns:
        tuple: (extracted_text, file_type)
        If file_type is 'image', text will be empty (handled by AI directly)
    """
    file_type = get_file_type(file_path)

    if file_type == "pdf":
        text = await parse_pdf(file_path, max_pages)
        return text, file_type

    elif file_type == "docx":
        text = await parse_docx(file_path)
        return text, file_type

    elif file_type == "doc_legacy":
        text = await parse_doc_legacy(file_path)
        return text, file_type

    elif file_type == "xlsx":
        text = await parse_xlsx(file_path)
        return text, file_type

    elif file_type == "image":
        # Images will be sent directly to Gemini's vision model
        return "", file_type

    else:
        logger.warning(f"Unsupported file type: {file_type} ({file_path})")
        return "", "unknown"
